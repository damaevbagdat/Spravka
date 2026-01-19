# -*- coding: utf-8 -*-
"""
Скрипт для генерации Коммерческого предложения ТОО "DES Строй"
"""

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from datetime import datetime
import os

# Для Word
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def num2text(n):
    """Конвертация числа в текст на русском (тенге)"""
    units = ['', 'один', 'два', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять']
    units_f = ['', 'одна', 'две', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять']
    teens = ['десять', 'одиннадцать', 'двенадцать', 'тринадцать', 'четырнадцать',
             'пятнадцать', 'шестнадцать', 'семнадцать', 'восемнадцать', 'девятнадцать']
    tens = ['', '', 'двадцать', 'тридцать', 'сорок', 'пятьдесят',
            'шестьдесят', 'семьдесят', 'восемьдесят', 'девяносто']
    hundreds = ['', 'сто', 'двести', 'триста', 'четыреста', 'пятьсот',
                'шестьсот', 'семьсот', 'восемьсот', 'девятьсот']

    if n == 0:
        return 'ноль тенге 00 тиын'

    n = int(n)
    result = []

    # Миллионы
    if n >= 1000000:
        m = n // 1000000
        n %= 1000000
        if m % 100 in range(11, 20):
            result.append(f"{hundreds[m // 100]} {teens[m % 100 - 10]} миллионов".strip())
        else:
            m_text = []
            if m >= 100:
                m_text.append(hundreds[m // 100])
            if (m % 100) >= 20:
                m_text.append(tens[(m % 100) // 10])
            if (m % 100) in range(10, 20):
                m_text.append(teens[m % 10])
            elif m % 10:
                m_text.append(units[m % 10])

            last = m % 10
            last_two = m % 100
            if last_two in range(11, 20):
                word = 'миллионов'
            elif last == 1:
                word = 'миллион'
            elif last in [2, 3, 4]:
                word = 'миллиона'
            else:
                word = 'миллионов'
            result.append(' '.join(m_text) + ' ' + word)

    # Тысячи
    if n >= 1000:
        t = n // 1000
        n %= 1000
        if t % 100 in range(11, 20):
            result.append(f"{hundreds[t // 100]} {teens[t % 100 - 10]} тысяч".strip())
        else:
            t_text = []
            if t >= 100:
                t_text.append(hundreds[t // 100])
            if (t % 100) >= 20:
                t_text.append(tens[(t % 100) // 10])
            if (t % 100) in range(10, 20):
                t_text.append(teens[t % 10])
            elif t % 10:
                t_text.append(units_f[t % 10])  # женский род для тысяч

            last = t % 10
            last_two = t % 100
            if last_two in range(11, 20):
                word = 'тысяч'
            elif last == 1:
                word = 'тысяча'
            elif last in [2, 3, 4]:
                word = 'тысячи'
            else:
                word = 'тысяч'
            result.append(' '.join(t_text) + ' ' + word)

    # Сотни, десятки, единицы
    if n > 0:
        n_text = []
        if n >= 100:
            n_text.append(hundreds[n // 100])
        if (n % 100) >= 20:
            n_text.append(tens[(n % 100) // 10])
        if (n % 100) in range(10, 20):
            n_text.append(teens[n % 10])
        elif n % 10:
            n_text.append(units[n % 10])
        result.append(' '.join(n_text))

    text = ' '.join(result).strip()
    # Капитализируем первую букву
    text = text[0].upper() + text[1:] if text else ''
    return f"{text} тенге 00 тиын"

def create_kp():
    wb = Workbook()
    ws = wb.active
    ws.title = "КП"

    # Настройка ширины колонок
    ws.column_dimensions['A'].width = 5
    ws.column_dimensions['B'].width = 45
    ws.column_dimensions['C'].width = 12
    ws.column_dimensions['D'].width = 8
    ws.column_dimensions['E'].width = 12
    ws.column_dimensions['F'].width = 14

    # Стили
    header_font = Font(bold=True, size=14)
    title_font = Font(bold=True, size=12)
    normal_font = Font(size=10)
    small_font = Font(size=9)

    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
    right_align = Alignment(horizontal='right', vertical='center')

    header_fill = PatternFill(start_color='E6E6E6', end_color='E6E6E6', fill_type='solid')

    # === ШАПКА КОМПАНИИ ===
    row = 1

    # Название компании
    ws.merge_cells('A1:F1')
    ws['A1'] = 'Товарищество с ограниченной ответственностью «DES Строй»'
    ws['A1'].font = Font(bold=True, size=14)
    ws['A1'].alignment = center_align

    row = 2
    ws.merge_cells('A2:F2')
    ws['A2'] = 'БИН: 240840038562'
    ws['A2'].font = normal_font
    ws['A2'].alignment = center_align

    row = 3
    ws.merge_cells('A3:F3')
    ws['A3'] = 'Адрес: Казахстан, Бостандыкский район, МИКРОРАЙОН АЛМАГУЛЬ, дом 27, кв/офис 1'
    ws['A3'].font = small_font
    ws['A3'].alignment = center_align

    row = 4
    ws.merge_cells('A4:F4')
    ws['A4'] = 'Тел.: +7 701 795 2575, +7 777 822 9141 | E-mail: des-stroy@mail.ru'
    ws['A4'].font = small_font
    ws['A4'].alignment = center_align

    row = 5
    ws.merge_cells('A5:F5')
    ws['A5'] = 'ИИК: KZ82 722S 0000 4248 1068 | Банк: АО «Kaspi Bank» | БИК: CASPKZKA'
    ws['A5'].font = small_font
    ws['A5'].alignment = center_align

    # Пустая строка
    row = 6

    # === ЗАГОЛОВОК КП ===
    row = 7
    ws.merge_cells('A7:F7')
    ws['A7'] = 'КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ'
    ws['A7'].font = Font(bold=True, size=16)
    ws['A7'].alignment = center_align

    # Дата
    row = 8
    ws.merge_cells('A8:F8')
    today = datetime.now().strftime('%d.%m.%Y')
    ws['A8'] = f'от {today} г.'
    ws['A8'].font = normal_font
    ws['A8'].alignment = center_align

    # Пустая строка
    row = 9

    # === ОБРАЩЕНИЕ ===
    row = 10
    ws.merge_cells('A10:F10')
    ws['A10'] = 'Уважаемые господа!'
    ws['A10'].font = Font(bold=True, size=11)
    ws['A10'].alignment = left_align

    row = 11
    ws.merge_cells('A11:F11')
    ws['A11'] = 'ТОО «DES Строй» предлагает Вам следующие товары/услуги:'
    ws['A11'].font = normal_font
    ws['A11'].alignment = left_align

    # Пустая строка
    row = 12

    # === ТАБЛИЦА ТОВАРОВ ===
    # Заголовки таблицы
    row = 13
    headers = ['№', 'Наименование', 'Кол-во', 'Ед.', 'Цена, тг', 'Сумма, тг']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=row, column=col, value=header)
        cell.font = Font(bold=True, size=10)
        cell.alignment = center_align
        cell.border = thin_border
        cell.fill = header_fill

    # Данные товаров
    items = [
        # Первый счет (сантехника)
        ('труба пвх ф50/250 (3.2) Deniz', 15, 'шт', 650.00),
        ('отвод пвх ф50*45 Deniz', 20, 'шт', 380.00),
        ('отвод пвх ф50*90 Deniz', 10, 'шт', 360.00),
        ('тройник пвх ф50*45 Deniz', 5, 'шт', 650.00),
        ('клипсы ф50 Deniz', 10, 'шт', 250.00),
        ('муфта пвх ф50 Deniz', 6, 'шт', 320.00),
        ('Бел труба стекло 20', 40, 'м', 530.00),
        ('Бел отвод 20', 30, 'шт', 65.00),
        ('Бел муфта 20', 10, 'шт', 55.00),
        ('Бел переход н 20x15', 12, 'шт', 650.00),
        ('Бел клипса 20', 15, 'шт', 50.00),
        ('муфта вставная 110 (восстановитель раструба) RTP', 1, 'шт', 7900.00),
        # Второй счет (стройматериалы Леруа Мерлен)
        ('Тумба Vigo Milk 60 56.6x43.4x70 см', 5, 'шт', 39790.00),
        ('Клей AlinEX «Set 301», 25 кг', 2, 'шт', 3180.00),
        ('Затирка цементная Axton A100 цвет серый 2 кг', 4, 'шт', 1570.00),
        ('Крестики для кафельной плитки Стройбат 2 мм 100 шт.', 10, 'шт', 245.00),
        ('Шпатель Dominus Профи 350 мм, нержавеющая сталь', 1, 'шт', 3510.00),
        ('Плитка настенная Шахтинская Плитка Моца 20x30 см 1.44 м² глянцевая цвет бежевый', 11, 'шт', 4550.00),
        ('Раковина Cersanit Colour керамика 60 см цвет белый', 5, 'шт', 38880.00),
        ('Смеситель для раковины Глория однорычажный', 5, 'шт', 11930.00),
        ('Сифон для раковины Equation с выпуском 32 мм', 5, 'шт', 2650.00),
        ('Крепеж РВК ф50 мм полипропилен для трубы с фиксатором', 10, 'шт', 125.00),
        # Дополнительные товары
        ('Мыльница', 5, 'шт', 2000.00),
        ('Салфетница', 5, 'шт', 2000.00),
    ]

    total = 0
    for i, (name, qty, unit, price) in enumerate(items, 1):
        row = 13 + i
        amount = qty * price
        total += amount

        # №
        cell = ws.cell(row=row, column=1, value=i)
        cell.font = normal_font
        cell.alignment = center_align
        cell.border = thin_border

        # Наименование
        cell = ws.cell(row=row, column=2, value=name)
        cell.font = normal_font
        cell.alignment = left_align
        cell.border = thin_border

        # Кол-во
        cell = ws.cell(row=row, column=3, value=qty)
        cell.font = normal_font
        cell.alignment = center_align
        cell.border = thin_border

        # Ед.
        cell = ws.cell(row=row, column=4, value=unit)
        cell.font = normal_font
        cell.alignment = center_align
        cell.border = thin_border

        # Цена
        cell = ws.cell(row=row, column=5, value=price)
        cell.font = normal_font
        cell.alignment = right_align
        cell.border = thin_border
        cell.number_format = '#,##0.00'

        # Сумма
        cell = ws.cell(row=row, column=6, value=amount)
        cell.font = normal_font
        cell.alignment = right_align
        cell.border = thin_border
        cell.number_format = '#,##0.00'

    # Строка ИТОГО
    row = 13 + len(items) + 1
    ws.merge_cells(f'A{row}:E{row}')
    cell = ws.cell(row=row, column=1, value='ИТОГО:')
    cell.font = Font(bold=True, size=11)
    cell.alignment = right_align
    cell.border = thin_border

    # Применяем границы к объединенным ячейкам
    for col in range(2, 6):
        ws.cell(row=row, column=col).border = thin_border

    cell = ws.cell(row=row, column=6, value=total)
    cell.font = Font(bold=True, size=11)
    cell.alignment = right_align
    cell.border = thin_border
    cell.number_format = '#,##0.00'

    # Сумма прописью
    row += 2
    ws.merge_cells(f'A{row}:F{row}')
    ws[f'A{row}'] = f'Всего наименований: {len(items)}, на сумму {total:,.2f} KZT'.replace(',', ' ')
    ws[f'A{row}'].font = Font(bold=True, size=10)
    ws[f'A{row}'].alignment = left_align

    row += 1
    ws.merge_cells(f'A{row}:F{row}')
    ws[f'A{row}'] = f'Всего к оплате: Шестьдесят восемь тысяч семьсот семьдесят тенге 00 тиын'
    ws[f'A{row}'].font = Font(bold=True, size=10)
    ws[f'A{row}'].alignment = left_align

    # === УСЛОВИЯ ===
    row += 2
    ws.merge_cells(f'A{row}:F{row}')
    ws[f'A{row}'] = 'Условия:'
    ws[f'A{row}'].font = Font(bold=True, size=11)

    row += 1
    ws.merge_cells(f'A{row}:F{row}')
    ws[f'A{row}'] = '• Срок действия предложения: 3 дня'
    ws[f'A{row}'].font = normal_font

    row += 1
    ws.merge_cells(f'A{row}:F{row}')
    ws[f'A{row}'] = '• Условия оплаты: 100% предоплата'
    ws[f'A{row}'].font = normal_font

    row += 1
    ws.merge_cells(f'A{row}:F{row}')
    ws[f'A{row}'] = '• Срок поставки: по согласованию'
    ws[f'A{row}'].font = normal_font

    # === ПОДПИСЬ ===
    row += 3
    ws.merge_cells(f'A{row}:F{row}')
    ws[f'A{row}'] = 'С уважением,'
    ws[f'A{row}'].font = normal_font

    row += 1
    ws.merge_cells(f'A{row}:F{row}')
    ws[f'A{row}'] = 'Директор ТОО «DES Строй»'
    ws[f'A{row}'].font = normal_font

    row += 1
    ws.merge_cells(f'A{row}:F{row}')
    ws[f'A{row}'] = 'Нысамбай Ердәулет Нұрманұлы _________________'
    ws[f'A{row}'].font = Font(bold=True, size=11)

    # Настройка печати
    ws.print_title_rows = '1:6'
    ws.page_setup.orientation = 'portrait'
    ws.page_setup.fitToPage = True
    ws.page_margins.left = 0.5
    ws.page_margins.right = 0.5
    ws.page_margins.top = 0.5
    ws.page_margins.bottom = 0.5

    # Сохранение
    output_path = r'c:\Users\damae\OneDrive\Документы\GitHub\Справка о задолженности\КП_DES_Строй.xlsx'
    wb.save(output_path)
    print(f'КП Excel сохранено: {output_path}')
    return output_path


def set_cell_border(cell, **kwargs):
    """Установка границ ячейки таблицы Word"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def create_kp_word():
    """Создание КП в формате Word"""
    doc = Document()

    # Настройка полей страницы
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)

    # === ШАПКА КОМПАНИИ ===
    # Название компании
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Товарищество с ограниченной ответственностью «DES Строй»')
    run.bold = True
    run.font.size = Pt(14)

    # БИН
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('БИН: 240840038562')
    run.font.size = Pt(10)

    # Адрес
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Адрес: Казахстан, Бостандыкский район, МИКРОРАЙОН АЛМАГУЛЬ, дом 27, кв/офис 1')
    run.font.size = Pt(9)

    # Контакты
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Тел.: +7 701 795 2575, +7 777 822 9141 | E-mail: des-stroy@mail.ru')
    run.font.size = Pt(9)

    # Банковские реквизиты
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ИИК: KZ82 722S 0000 4248 1068 | Банк: АО «Kaspi Bank» | БИК: CASPKZKA')
    run.font.size = Pt(9)

    # Пустая строка
    doc.add_paragraph()

    # === ЗАГОЛОВОК КП ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ')
    run.bold = True
    run.font.size = Pt(16)

    # Дата
    today = datetime.now().strftime('%d.%m.%Y')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'от {today} г.')
    run.font.size = Pt(10)

    # Пустая строка
    doc.add_paragraph()

    # === ОБРАЩЕНИЕ ===
    p = doc.add_paragraph()
    run = p.add_run('Уважаемые господа!')
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run('ТОО «DES Строй» предлагает Вам следующие товары/услуги:')
    run.font.size = Pt(10)

    # === ДАННЫЕ ТОВАРОВ ===
    items = [
        # Первый счет (сантехника)
        ('труба пвх ф50/250 (3.2) Deniz', 15, 'шт', 650.00),
        ('отвод пвх ф50*45 Deniz', 20, 'шт', 380.00),
        ('отвод пвх ф50*90 Deniz', 10, 'шт', 360.00),
        ('тройник пвх ф50*45 Deniz', 5, 'шт', 650.00),
        ('клипсы ф50 Deniz', 10, 'шт', 250.00),
        ('муфта пвх ф50 Deniz', 6, 'шт', 320.00),
        ('Бел труба стекло 20', 40, 'м', 530.00),
        ('Бел отвод 20', 30, 'шт', 65.00),
        ('Бел муфта 20', 10, 'шт', 55.00),
        ('Бел переход н 20x15', 12, 'шт', 650.00),
        ('Бел клипса 20', 15, 'шт', 50.00),
        ('муфта вставная 110 (восстановитель раструба) RTP', 1, 'шт', 7900.00),
        # Второй счет (стройматериалы Леруа Мерлен)
        ('Тумба Vigo Milk 60 56.6x43.4x70 см', 5, 'шт', 39790.00),
        ('Клей AlinEX «Set 301», 25 кг', 2, 'шт', 3180.00),
        ('Затирка цементная Axton A100 цвет серый 2 кг', 4, 'шт', 1570.00),
        ('Крестики для кафельной плитки Стройбат 2 мм 100 шт.', 10, 'шт', 245.00),
        ('Шпатель Dominus Профи 350 мм, нержавеющая сталь', 1, 'шт', 3510.00),
        ('Плитка настенная Шахтинская Плитка Моца 20x30 см 1.44 м² глянцевая цвет бежевый', 11, 'шт', 4550.00),
        ('Раковина Cersanit Colour керамика 60 см цвет белый', 5, 'шт', 38880.00),
        ('Смеситель для раковины Глория однорычажный', 5, 'шт', 11930.00),
        ('Сифон для раковины Equation с выпуском 32 мм', 5, 'шт', 2650.00),
        ('Крепеж РВК ф50 мм полипропилен для трубы с фиксатором', 10, 'шт', 125.00),
        # Дополнительные товары
        ('Мыльница', 5, 'шт', 2000.00),
        ('Салфетница', 5, 'шт', 2000.00),
    ]

    # === ТАБЛИЦА ТОВАРОВ ===
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовки
    headers = ['№', 'Наименование', 'Кол-во', 'Ед.', 'Цена, тг', 'Сумма, тг']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Границы
        set_cell_border(header_cells[i],
            top={"sz": 8, "val": "single", "color": "000000"},
            bottom={"sz": 8, "val": "single", "color": "000000"},
            left={"sz": 8, "val": "single", "color": "000000"},
            right={"sz": 8, "val": "single", "color": "000000"})

    # Данные
    total = 0
    for i, (name, qty, unit, price) in enumerate(items, 1):
        amount = qty * price
        total += amount

        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = name
        row[2].text = str(qty)
        row[3].text = unit
        row[4].text = f'{price:,.2f}'.replace(',', ' ')
        row[5].text = f'{amount:,.2f}'.replace(',', ' ')

        # Выравнивание и границы
        for j, cell in enumerate(row):
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if j in [0, 2, 3]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif j in [4, 5]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_cell_border(cell,
                top={"sz": 4, "val": "single", "color": "000000"},
                bottom={"sz": 4, "val": "single", "color": "000000"},
                left={"sz": 4, "val": "single", "color": "000000"},
                right={"sz": 4, "val": "single", "color": "000000"})

    # Строка ИТОГО
    row = table.add_row().cells
    row[0].merge(row[4])
    row[0].text = 'ИТОГО:'
    row[0].paragraphs[0].runs[0].bold = True
    row[0].paragraphs[0].runs[0].font.size = Pt(11)
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row[5].text = f'{total:,.2f}'.replace(',', ' ')
    row[5].paragraphs[0].runs[0].bold = True
    row[5].paragraphs[0].runs[0].font.size = Pt(11)
    row[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for cell in [row[0], row[5]]:
        set_cell_border(cell,
            top={"sz": 8, "val": "single", "color": "000000"},
            bottom={"sz": 8, "val": "single", "color": "000000"},
            left={"sz": 8, "val": "single", "color": "000000"},
            right={"sz": 8, "val": "single", "color": "000000"})

    # Настройка ширины колонок
    widths = [Cm(1), Cm(8), Cm(1.5), Cm(1), Cm(2.5), Cm(3)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

    # Пустая строка
    doc.add_paragraph()

    # Сумма прописью
    p = doc.add_paragraph()
    run = p.add_run(f'Всего наименований: {len(items)}, на сумму {total:,.2f} KZT'.replace(',', ' '))
    run.bold = True
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run(f'Всего к оплате: {num2text(total)}')
    run.bold = True
    run.font.size = Pt(10)

    # Пустая строка
    doc.add_paragraph()

    # === УСЛОВИЯ ===
    p = doc.add_paragraph()
    run = p.add_run('Условия:')
    run.bold = True
    run.font.size = Pt(11)

    conditions = [
        '• Срок действия предложения: 3 дня',
        '• Условия оплаты: 100% предоплата',
        '• Срок поставки: 21 рабочий день после оплаты',
        '• Общая стоимость товаров включает НДС'
    ]
    for cond in conditions:
        p = doc.add_paragraph()
        run = p.add_run(cond)
        run.font.size = Pt(10)

    # Пустая строка
    doc.add_paragraph()
    doc.add_paragraph()

    # === ПОДПИСЬ ===
    p = doc.add_paragraph()
    run = p.add_run('С уважением,')
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run('Директор ТОО «DES Строй»')
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run('Нысамбай Ердәулет Нұрманұлы _________________')
    run.bold = True
    run.font.size = Pt(11)

    # Сохранение
    output_path = r'c:\Users\damae\OneDrive\Документы\GitHub\Справка о задолженности\КП_DES_Строй.docx'
    doc.save(output_path)
    print(f'КП Word сохранено: {output_path}')
    return output_path


if __name__ == '__main__':
    create_kp()  # Excel версия
    create_kp_word()  # Word версия
